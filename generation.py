# generation.py

import re
import pathlib
from datetime import datetime
import pandas as pd
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from fastapi import HTTPException

# Symboles pour Oui/Non
CHECKED, UNCHECKED = "☑", "☐"

# Mapping des colonnes Excel vers balises du template
MAP = {
    "Nom du candidat":            "NOM",
    "Date prise de références":   "DATE_PR",
    "Statut":                     "STATUT",
    "Département":                "POSTE",
    "Candidature retenue ou non": "RETENU",
    "Date entrevue personne":     "DATE_ENT_PERS",
    "Candidat rejoint oui/non":   "REJOINT",
    "Date entrevue téléphonique": "DATE_ENT_TEL",
}

def _iter_paragraphs_and_cells(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield cell
                for p in cell.paragraphs:
                    yield p

def _rewrite_runs(obj: Paragraph | _Cell, new_text: str):
    if isinstance(obj, _Cell):
        for p in obj.paragraphs:
            _rewrite_runs(p, new_text)
        return
    for r in obj.runs:
        r.text = ""
    if obj.runs:
        obj.runs[0].text = new_text
    else:
        obj.add_run(new_text)

def replace_text(obj: Paragraph | _Cell, old: str, new: str):
    if isinstance(obj, _Cell):
        for p in obj.paragraphs:
            replace_text(p, old, new)
        return
    if old in obj.text:
        _rewrite_runs(obj, obj.text.replace(old, new))

def mark_choice(obj: Paragraph | _Cell, keyword: str, yes_selected: bool):
    if isinstance(obj, _Cell):
        for p in obj.paragraphs:
            if mark_choice(p, keyword, yes_selected):
                return True
        return False
    if keyword not in obj.text:
        return False
    def repl(m):
        word = m.group(0)
        if word.lower().startswith("oui"):
            return f"{CHECKED} Oui" if yes_selected else f"{UNCHECKED} Oui"
        else:
            return f"{UNCHECKED} Non" if yes_selected else f"{CHECKED} Non"
    new = re.sub(r"\b(Oui|Non)\b", repl, obj.text, count=2, flags=re.I)
    _rewrite_runs(obj, new)
    return True

def run_reporter(excel_fp: pathlib.Path, tpl_fp: pathlib.Path, out_dir: pathlib.Path):
    # Lecture avec la première ligne comme entêtes
    df = pd.read_excel(excel_fp, dtype=str)
    # Vérifie la présence de la colonne principale
    if "Nom du candidat" not in df.columns:
        raise HTTPException(400, 'Le fichier Excel doit contenir la colonne "Nom du candidat".')
    # Netoyer les NaN et forcer en str
    df = df.fillna("").astype(str)

    stamp = datetime.now().strftime("%Y%m%d")
    out_dir.mkdir(exist_ok=True)

    for _, row in df.iterrows():
        nom = row.get("Nom du candidat", "").strip()
        if not nom:
            continue

        vals = {}
        for col, key in MAP.items():
            raw = row.get(col, "")
            v = raw.strip()
            v = re.sub(r"\s+00:00:00$", "", v)
            vals[key] = v if v and v.lower() != "nan" else "N/A"

        retenu = vals["RETENU"].lower().startswith("oui")
        rejoint = vals["REJOINT"].lower().startswith("oui")

        doc = Document(tpl_fp)
        for obj in _iter_paragraphs_and_cells(doc):
            replace_text(obj, "<<NOM>>", nom)
            replace_text(obj, "<<DATE_PR>>", vals["DATE_PR"])
            replace_text(obj, "<<STATUT>>", vals["STATUT"])
            replace_text(obj, "<<POSTE>>", vals["POSTE"])
            replace_text(obj, "<<DATE_ENT_PERS>>", vals["DATE_ENT_PERS"])
            replace_text(obj, "<<DATE_ENT_TEL>>", vals["DATE_ENT_TEL"])
            msg = vals["DATE_ENT_TEL"] if rejoint else ""
            replace_text(obj, "<<MSG_DATE>>", msg)
            mark_choice(obj, "Candidat retenu", retenu)
            mark_choice(obj, "Message laissé", rejoint)

        out_name = f"Form_{nom.replace(' ', '_')}_{stamp}.docx"
        doc.save(out_dir / out_name)
